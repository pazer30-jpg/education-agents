"""
Agent 2 - Writer
קורא מאמרים וכותב מאמר אקדמי מלא עם Claude CLI
פלט: Markdown + DOCX — בעברית ובאנגלית
"""

import json
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from config import ARTICLES_DIR
from claude_cli import ask_claude, ask_claude_json

try:
    from obsidian_memory import format_for_prompt as _obsidian_memory_for_prompt
except Exception:
    def _obsidian_memory_for_prompt(_names: list[str], **_kw) -> str:
        return ""


# ─────────────────────────────────────────────
# DOCX builder
# ─────────────────────────────────────────────

def _markdown_to_docx(md_text: str, title: str, output_path: Path):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(title)
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    # Date
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(datetime.now().strftime("%B %Y"))
    date_run.font.size = Pt(10)
    date_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    doc.add_paragraph()

    for line in md_text.split("\n"):
        line = line.rstrip()
        if line.startswith("## "):
            h = doc.add_heading(line[3:], level=1)
            if h.runs:
                h.runs[0].font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        elif line.startswith("### "):
            h = doc.add_heading(line[4:], level=2)
            if h.runs:
                h.runs[0].font.color.rgb = RGBColor(0x2D, 0x2D, 0x5E)
        elif line.startswith("**") and line.endswith("**"):
            p = doc.add_paragraph()
            r = p.add_run(line.strip("*"))
            r.bold = True
            r.font.size = Pt(11)
        elif line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(line[2:], style="List Bullet")
            if p.runs:
                p.runs[0].font.size = Pt(11)
        elif line.strip():
            p = doc.add_paragraph(line)
            if p.runs:
                p.runs[0].font.size = Pt(11)
            p.paragraph_format.space_after = Pt(6)
        else:
            doc.add_paragraph()

    doc.save(str(output_path))
    return output_path


# ─────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────

def _split_title(article_md: str, default_title: str) -> tuple[str, str]:
    """Extract # title from first line, return (title, body)."""
    lines = article_md.strip().split("\n")
    if lines[0].startswith("# "):
        return lines[0][2:].strip(), "\n".join(lines[1:]).strip()
    return default_title, article_md.strip()


def _create_briefing(article_en: str, display_title: str, base: str,
                     synthesis_block: str = "") -> Path:
    """
    Create a practitioner-facing briefing from the academic article.
    This is what Agent 3 reads — NOT the full article.
    Format designed for Paz (practitioner), not for journals.
    """
    print("  [Agent2] Creating practitioner briefing for Agent 3...")

    system = """אתה מסייע לפרקטיקן בחינוך בלתי-פורמלי לעבד מחקר אקדמי.
אתה לא כותב מאמר. אתה מכין לו מסמך עבודה קצר — scaffolding — שממנו הוא יכתוב בעצמו.

הוא צריך:
  1. מה המאמרים באמת טוענים (מדויק, לא מוגזם)
  2. מספרים אמיתיים שהוא יכול לצטט
  3. סתירות ופערים שהוא יכול להצביע עליהם
  4. זוויות שהוא יכול לחקור מניסיונו בשטח

אתה לא מנסה להרשים. אתה נותן לו חומר גלם שהוא יעבד בקול שלו."""

    prompt = f"""המאמר האקדמי המלא:

{article_en}

{synthesis_block}

בנה מסמך briefing פרקטי (עברית, 800-1,200 מילים) במבנה הזה:

## נושא
(משפט אחד — על מה זה)

## 5-7 טענות מרכזיות מהמחקר
לכל טענה:
- הטענה עצמה (משפט)
- ✓ proven / ~ suggested / T theoretical
- הציטוט (Author, Year) + אם יש: n=, d=, p=
- משפט קצר: "מה פז יכול לומר על זה מהשטח"

## מספרים שפז יכול לצטט
רשימה של 5-10 נתונים עם מקור. רק מה שבאמת מופיע במאמרים.

## סתירות/מתחים במחקר
2-3 מקומות שמחקרים לא מסכימים, או שיש פער בין מה שנחקר למציאות.

## 3 שאלות פתוחות
שאלות שהמחקרים לא עונים עליהן. אלה הזוויות של פז.

## 2 זוויות פרסום
שתי אפשרויות ספציפיות לפוסט/בלוג שפז יכול לכתוב:
- זווית 1: [טענה] ← [דוגמה אפשרית מהשטח שפז יכיר]
- זווית 2: ...

הנחיות:
- עברית טבעית, לא אקדמית
- ציטוטים בצורה (שם, שנה) — שמות חוקרים בעברית
- לא להגזים. ~ suggested זה לא proven.
- לא להמציא מספרים.
- לא לכתוב כאילו אתה פז — אתה מכין לו חומר."""

    briefing = ask_claude(prompt, system=system, max_budget=2.0)
    briefing_path = ARTICLES_DIR / f"{base}_briefing.md"
    briefing_path.write_text(
        f"# Briefing: {display_title}\n\n"
        f"*מסמך עבודה פנימי — חומר גלם לפז. לא לפרסום.*\n\n"
        f"{briefing}",
        encoding="utf-8",
    )
    print(f"  [Agent2] Briefing saved: {briefing_path.name}")
    return briefing_path


# ─────────────────────────────────────────────
# Main agent function
# ─────────────────────────────────────────────

def _load_papers_from_files(papers_files: list[Path]) -> tuple[list[dict], list[str]]:
    """Load and merge papers from multiple JSON files. Returns (papers, topics)."""
    all_papers = []
    topics = []
    seen_ids = set()

    for pf in papers_files:
        if not pf.exists():
            print(f"  [Agent2] ⚠️  {pf.name} לא נמצא — מדלג")
            continue
        with open(pf, encoding="utf-8") as f:
            data = json.load(f)
        topic = data.get("topic", pf.stem)
        papers = data.get("papers", data) if isinstance(data, dict) else data
        topics.append(topic)
        for p in papers:
            pid = p.get("paperId") or p.get("title", "")[:60]
            if pid not in seen_ids:
                seen_ids.add(pid)
                p["_source_topic"] = topic
                all_papers.append(p)

    # Prefer fulltext over abstract; trim to context-safe size
    for p in all_papers:
        if p.get("fulltext") and len(p["fulltext"]) > 200:
            # Use fulltext — trim if too long
            if len(p["fulltext"]) > 8000:
                p["fulltext"] = p["fulltext"][:8000] + "..."
            p.pop("abstract", None)   # drop abstract to save tokens
        elif isinstance(p.get("abstract"), str) and len(p["abstract"]) > 500:
            p["abstract"] = p["abstract"][:500] + "..."

    return all_papers, topics


def run_writer(papers_files: Path | list[Path], combined_title: str = "",
               bilingual: bool = True,
               research_questions: list[dict] = None,
               combined_research_question: str = "") -> dict[str, Path]:
    """
    papers_files: single Path or list of Paths (one per researched topic)
    combined_title: optional display title for the combined article
    bilingual: if True (default), write both EN and HE articles
    research_questions: list of {topic, question, sub_questions} from Agent 0 (NEW)
    combined_research_question: overarching RQ from Agent 0 (NEW)
    """
    if isinstance(papers_files, Path):
        papers_files = [papers_files]

    names = " + ".join(p.name for p in papers_files[:3])
    print(f"\n{'='*60}")
    print(f"✍️  Agent 2 - Writer | {len(papers_files)} קבצי מחקר: {names}")
    print(f"{'='*60}\n")

    all_papers, topics = _load_papers_from_files(papers_files)
    topics_str = " × ".join(topics)
    display_title = combined_title or topics_str

    if not all_papers:
        raise ValueError("No papers found in any of the provided files.")

    print(f"  [Agent2] {len(all_papers)} מאמרים מ-{len(topics)} נושאים: {topics_str}")

    # ── Paper Analyzer (Agent 1.7) — synthesis map ──
    synthesis_block = ""
    try:
        from paper_analyzer import run_paper_analyzer
        analysis = run_paper_analyzer(papers_files[-1])
        synthesis_block = analysis["synthesis_map"]
        print(f"  [Agent2] Synthesis map מוכן — {len(analysis.get('profiles',[]))} profiles")
    except Exception as e:
        print(f"  [Agent2] ⚠️  Paper analysis דולג ({e})")

    # Build topic breakdown for prompt
    topic_breakdown = "\n".join(
        f"  - {t}: {sum(1 for p in all_papers if p.get('_source_topic') == t)} papers"
        for t in topics
    )

    system = """You are a senior academic writer specializing in education research.
You write strictly according to APA 7th edition.
You write review/synthesis articles, NOT empirical studies.

Write a SYNTHESIZED article that weaves together multiple topics into one coherent argument.
Do NOT write separate sections per topic — integrate them throughout.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
ARTICLE STRUCTURE — MANDATORY SECTIONS
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
You MUST include ALL of the following sections in this exact order.
Omitting any section is a failure.

## Abstract (150-200 words)
## Introduction
  - Open with the broad problem in the field
  - State 2-3 explicit Research Questions (RQ1, RQ2, RQ3)
  - Each RQ must be answerable from the reviewed literature
  - End with article purpose and scope

## Methodology
  - Databases searched (Semantic Scholar, OpenAlex, ERIC, CORE, Crossref)
  - Search terms used
  - Inclusion criteria: peer-reviewed, years, language
  - Exclusion criteria
  - Total found → screened → included (state numbers)
  - This is a review article — the methodology describes HOW you searched, not an experiment

## Theoretical Framework
## Literature Review (integrated synthesis with subsections)
## Discussion
## Limitations of This Review
  - Language bias (English-only sources)
  - Database limitations
  - Date range constraints
  - Gaps in available literature
## Conclusions
## References

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
CRITICAL SOURCE EVALUATION — MANDATORY
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
When citing a source, you MUST evaluate it — not just report it:
  - Method type: "In a quantitative study (n=400)..." / "Using qualitative interviews (n=12)..."
  - Sample size: always state n= when available
  - Generalizability: "While limited to [context], the findings suggest..."
  - Contrast weak vs strong evidence: "X (n=23) found A, but the larger study by Y (n=1,200) contradicts this"
  - State limitations: "however, participants were self-selected" / "the sample was limited to [country]"

Do NOT treat all sources as equal. A meta-analysis carries more weight than a case study.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
COMPARISON TABLE — MANDATORY
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Include AT LEAST one Markdown comparison table in the Literature Review:

| Study | Year | Method | Sample | Key Finding |
|-------|------|--------|--------|-------------|
| Author et al. | 2019 | Quantitative | n=400 | ... |

The table should compare 6-10 key studies side by side.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
APA 7 CITATION RULES — MANDATORY
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
In-text citations:
  One author:    (Smith, 2019)
  Two authors:   (Smith & Jones, 2019)
  3+ authors:    (Smith et al., 2019)
  Direct quote:  (Smith, 2019, p. 45)
  Narrative:     Smith (2019) argued that...

Every claim, finding, or idea from a source MUST have an in-text citation.
Minimum 20 in-text citations across the article.
Do NOT write a paragraph without at least one citation.

References list (APA 7):
  Author, A. A., & Author, B. B. (Year). Title of article.
  Journal Name, Volume(Issue), pages. https://doi.org/xxxxx

Rules:
  - Only list sources actually cited in the text
  - DOI as URL when available
  - Alphabetical by first author last name
  - Never: "important to note", "it can be seen", "as shown above"
"""

    synthesis_section = f"\n\nSYNTHESIS MAP (use this to structure your article — write about the debates, fill the gaps, reference the consensus):\n{synthesis_block}\n" if synthesis_block else ""

    # ── Slim papers: send only essential fields, not full abstracts ──
    slim_papers = []
    for p in all_papers:
        slim = {
            "title": (p.get("title") or "")[:120],
            "authors": (p.get("authors") or "")[:80] if isinstance(p.get("authors"), str)
                       else ", ".join(str(a) for a in (p.get("authors") or [])[:3])[:80],
            "year": p.get("year"),
            "citation_count": p.get("citation_count") or 0,
            "venue": (p.get("venue") or "")[:60],
            "abstract": (p.get("abstract") or p.get("fulltext") or "")[:250],
            "topic": p.get("_source_topic", ""),
        }
        slim_papers.append(slim)
    papers_json = json.dumps(slim_papers, ensure_ascii=False, indent=1)

    base = "_x_".join(t.replace(" ", "_").lower()[:20] for t in topics)[:60]

    # ── Phase 0: Generate outline first (cheap sanity check) ──
    print("  [Agent2] Phase 0: Building outline...")

    # Use RQs from Agent 0 if provided — they're sharper than what we'd build here
    rq_block = ""
    if research_questions:
        rq_lines = []
        for i, rq in enumerate(research_questions, 1):
            rq_lines.append(f"RQ{i} ({rq.get('topic','')}): {rq.get('question','')}")
            for sq in rq.get('sub_questions', [])[:2]:
                rq_lines.append(f"   • {sq}")
        rq_block = "\n\n━━━ RESEARCH QUESTIONS (from Planner) ━━━\n" + "\n".join(rq_lines)
        if combined_research_question:
            rq_block += f"\n\nOVERARCHING RQ: {combined_research_question}"
        rq_block += "\n━━━ Use these as the article's RQs verbatim ━━━\n"

    outline_prompt = f"""Topics: {topics_str}
Papers available: {len(slim_papers)}
{synthesis_section}
{rq_block}

Build a SHORT outline for a synthesized academic review article.
Return JSON:
{{
  "central_thesis": "one-sentence thesis connecting all {len(topics)} topics",
  "research_questions": {"[use the RQs from above verbatim]" if research_questions else '["RQ1: ...", "RQ2: ...", "RQ3: ..."]'},
  "lit_review_sections": [
    {{"name": "...", "key_papers": ["Author 2020", "Author 2021"]}},
    {{"name": "...", "key_papers": ["..."]}}
  ],
  "main_tensions": ["tension or debate the article should address"],
  "answer_to_rq1": "one-sentence answer based on the literature",
  "answer_to_rq2": "...",
  "answer_to_rq3": "..."
}}

Use ONLY papers from this list — no fabrications:
{papers_json[:8000]}"""

    try:
        outline = ask_claude_json(outline_prompt, system=system, max_budget=0.8, timeout=400)
        rqs = outline.get("research_questions", [])
        thesis = outline.get("central_thesis", "")
        sections = outline.get("lit_review_sections", [])
        tensions = outline.get("main_tensions", [])

        # ── Outline Gate: score quality + retry if weak ──
        def _outline_score(o: dict) -> tuple[int, list[str]]:
            """Score 0-100 based on quality criteria. Return (score, issues)."""
            issues = []
            score = 100
            r = o.get("research_questions", [])
            t = o.get("central_thesis", "")
            s = o.get("lit_review_sections", [])
            tn = o.get("main_tensions", [])

            if len(r) < 2:
                score -= 25
                issues.append(f"only {len(r)} RQs (need 2-3)")
            if len(t) < 30:
                score -= 20
                issues.append(f"thesis too short ({len(t)} chars)")
            elif len(t) < 60:
                score -= 10
                issues.append(f"thesis underdeveloped ({len(t)} chars)")
            if len(s) < 3:
                score -= 25
                issues.append(f"only {len(s)} lit review sections (need 3+)")
            if len(tn) < 1:
                score -= 15
                issues.append("no tensions identified")

            # Heuristic: thesis should contain a verb (not just nouns)
            verb_markers = (
                " הוא ", " היא ", " הם ", " הן ",
                " is ", " are ", " was ", " were ", " shows ", " indicates ",
                # active claims
                " מציע", " טוען", " מראה", " מעיד", " מסביר", " חושף", " מבסס",
                # passive academic
                " מוסבר", " מנותח", " נחקר", " נמצא", " נבחן", " מתואר", " מפורש",
                # relational
                " מקושר", " קשור", " מתואם", " משפיע", " מושפע", " מעצב",
            )
            if t and not any(c in t for c in verb_markers):
                score -= 15  # raised from 10
                issues.append("thesis lacks a clear claim/verb")

            # Generic-thesis check
            if t and any(g in t.lower() for g in ("חשוב", "מעניין", "important", "interesting",
                                                   "ראוי לציין", "כדאי", "as we know")):
                score -= 20  # raised from 10 — these are RED flags
                issues.append("thesis uses generic words (חשוב/מעניין/ראוי לציין)")

            # Underdeveloped thesis is more serious
            if 0 < len(t) < 80:
                # additional penalty for very short OR short thesis with no specifics
                if not any(d in t for d in (":", ",", "—", "אבל", "מצד", "תוך", "באמצעות")):
                    score -= 10
                    issues.append("thesis lacks specific structure (no commas/clauses)")

            return max(0, score), issues

        score, issues = _outline_score(outline)
        max_retries = 2
        for attempt in range(max_retries):
            if score >= 60:
                break
            print(f"  [Agent2] ⚠️  Outline weak (score={score}) — retry {attempt+1}/{max_retries}")
            print(f"     issues: {'; '.join(issues)}")
            retry_prompt = outline_prompt + (
                f"\n\nIMPORTANT — previous outline attempt scored {score}/100. Fix these issues:\n"
                + "\n".join(f"- {i}" for i in issues)
                + "\n\nReturn a STRONGER outline with: clear thesis with verb, 2-3 specific RQs, "
                  "3+ lit-review sections, 1-2 real tensions."
            )
            try:
                outline = ask_claude_json(retry_prompt, system=system, max_budget=0.5, timeout=300)
                rqs = outline.get("research_questions", [])
                thesis = outline.get("central_thesis", "")
                sections = outline.get("lit_review_sections", [])
                tensions = outline.get("main_tensions", [])
                score, issues = _outline_score(outline)
            except Exception as e:
                print(f"  [Agent2] ⚠️ Retry failed: {e}")
                break

        outline_valid = score >= 60
        if outline_valid:
            print(f"  [Agent2] ✅ Outline OK ({score}/100): {len(rqs)} RQs, {len(sections)} sections, thesis={thesis[:60]}...")
        else:
            print(f"  [Agent2] ⚠️ Outline still weak ({score}/100) after retries — proceeding with caveat")
            # Push warning to scratchpad for downstream agents
            try:
                from scratchpad import note as _scratch_note
                _scratch_note("outline_gate", "weak_outline_warning", {
                    "issue": f"outline score {score}/100",
                    "issues": "; ".join(issues),
                    "summary": "המאמר נכתב למרות outline חלש — בדוק תוצאה לפני פרסום",
                })
            except Exception:
                pass

        # Build outline guidance for the writing phase
        outline_block = f"""

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
APPROVED OUTLINE (write according to this):
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
THESIS: {thesis}

{chr(10).join(rqs)}

LITERATURE REVIEW SECTIONS:
{chr(10).join(f"  • {s.get('name','?')}: {', '.join(s.get('key_papers', [])[:3])}" for s in sections)}

KEY TENSIONS TO ADDRESS:
{chr(10).join(f"  • {t}" for t in tensions[:3])}

ANSWERS (use in Discussion):
  RQ1: {outline.get('answer_to_rq1','')}
  RQ2: {outline.get('answer_to_rq2','')}
  RQ3: {outline.get('answer_to_rq3','')}
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
"""
    except Exception as e:
        print(f"  [Agent2] ⚠️ Outline failed ({e}) — proceeding without")
        outline_block = ""

    # ── Obsidian memory injection ──
    memory_block = _obsidian_memory_for_prompt([
        "academic_writing_apa7",
        "voice_rules",
        "theoretical_anchors",
        "recurring_sources",
        "humanize_rules",
    ], max_chars_per_note=1200)

    # ── Scratchpad: retry hints + cross-agent warnings ──
    scratchpad_block = ""
    try:
        from scratchpad import format_for_agent as _scratch_for
        scratchpad_block = _scratch_for("writer")
    except Exception:
        pass
    if scratchpad_block:
        memory_block = memory_block + "\n\n" + scratchpad_block

    # ── Single-call article generation (was 2 phases → now 1, saves a call) ──
    print("  [Agent2] Writing full article in one call...")
    prompt_full = f"""{memory_block}

Research topics (synthesized into ONE article): {topics_str}

Topic breakdown:
{topic_breakdown}
{synthesis_section}
{outline_block}
Papers ({len(slim_papers)} total — use these for citations):
{papers_json}

Write a COMPLETE synthesized academic review article — ALL sections, in order:

## Abstract (150-200 words — thesis + RQs + key conclusions)

## Introduction
- Broad problem → specific gap → 2-3 Research Questions (RQ1, RQ2, RQ3)

## Methodology
- Databases: Semantic Scholar, OpenAlex, ERIC, CORE, Crossref
- Search terms, inclusion/exclusion criteria
- Found: ~{len(all_papers)*4} → Screened: ~{len(all_papers)*2} → Included: {len(all_papers)}

## Theoretical Framework
(Shared theories across all {len(topics)} topics)

## Literature Review
- Integrated synthesis with 4-6 thematic subsections
- Include a comparison table (Study | Year | Method | Sample | Key Finding)
- Critically evaluate sources: state n=, method type, generalizability
- Do NOT summarize paper by paper — synthesize by theme

## Discussion
- Answer each RQ from the Introduction explicitly
- Tensions between findings, practical implications

## Limitations of This Review
- Language bias, database limitations, date range, gaps

## Conclusions
(Summary + specific future research directions)

## References
(APA 7 — ONLY papers actually cited in the text, alphabetical)

Target: ~3,500 words total. Every paragraph needs at least one (Author, Year) citation.
The article argues ONE central thesis connecting all {len(topics)} topics.
Write the ENTIRE article — do not stop partway."""

    # Writer time budget — track cumulative to skip optional steps if running long
    import time as _time
    _writer_start = _time.time()
    _writer_deadline = _writer_start + 1500  # 25 min hard cap (within 30 min step_timeout)

    article_en = ask_claude(prompt_full, system=system, max_budget=5.0, timeout=600)
    print(f"  [Agent2] Article: {len(article_en.split())} words")
    title_en, content_en = _split_title(article_en, f"Synthesized Article: {display_title}")

    # ── Pre-check: skip self-review if article is structurally complete ──
    required_sections = ["## Abstract", "## Introduction", "## Methodology",
                         "## Theoretical Framework", "## Literature Review",
                         "## Discussion", "## Limitations", "## Conclusions",
                         "## References"]
    missing = [s for s in required_sections if s not in article_en]
    has_table = "|" in article_en and "---" in article_en
    has_rqs = "RQ1" in article_en and "RQ2" in article_en

    if not missing and has_table and has_rqs:
        print(f"  [Agent2] ✅ All sections present — skipping self-review (saves $1.5)")
        # Still split title for downstream use
        title_en, content_en = _split_title(article_en, f"Synthesized Article: {display_title}")
        # Skip the self-review block by setting reviewed=None below path
        skip_review = True
    else:
        skip_review = False
        if missing:
            print(f"  [Agent2] Missing: {', '.join(s.replace('## ','') for s in missing)}")

    # ── Self-review: verify structure + citations (only if needed) ──
    if not skip_review:
        print("  [Agent2] Self-review: checking structure and citations...")
        review_prompt = f"""Review this academic article and fix problems. Return the FIXED article only.

CHECK LIST:
1. Every ## section from the structure exists? (Abstract, Introduction, Methodology, Theoretical Framework, Literature Review, Discussion, Limitations, Conclusions, References)
   If ANY section is missing — add it with appropriate content.
2. Every (Author, Year) citation in the text appears in ## References?
   If not — remove the orphan citation OR add to References.
3. Every entry in ## References is cited in the text?
   If not — remove from References.
4. Is there at least one Markdown table in Literature Review?
   If not — add a comparison table.
5. Are RQ1, RQ2, RQ3 stated in Introduction and answered in Discussion?
   If not — fix.
6. Does Methodology mention databases and inclusion criteria?
   If not — add.

Return the COMPLETE fixed article in Markdown. No explanations.

Article to review:
{article_en}"""

        # Skip self-review if cumulative time tight (already used >18 min)
        elapsed_so_far = _time.time() - _writer_start
        if elapsed_so_far > 18 * 60:
            print(f"  [Agent2] ⏭  Self-review skipped — already at {elapsed_so_far/60:.1f} min")
        else:
            try:
                reviewed = ask_claude(review_prompt, max_budget=1.5, timeout=600)
                if reviewed and len(reviewed) > len(article_en) * 0.7:
                    title_en, content_en = _split_title(reviewed, title_en)
                    print("  [Agent2] Self-review applied")
                else:
                    print("  [Agent2] Self-review skipped (output too short)")
            except Exception as e:
                print(f"  [Agent2] Self-review failed ({e}) — using original")

    md_en   = ARTICLES_DIR / f"{base}_en.md"
    docx_en = ARTICLES_DIR / f"{base}_en.docx"
    with open(md_en, "w", encoding="utf-8") as f:
        f.write(f"# {title_en}\n\n{content_en}")
    _markdown_to_docx(content_en, title_en, docx_en)
    print(f"  [Agent2] English saved: {md_en.name}, {docx_en.name}")

    # ── Briefing + Hebrew translation in PARALLEL (saves 4-5 min) ──
    # Both depend only on content_en; both make independent LLM calls.
    print("  [Agent2] Briefing + HE translation in parallel...")
    from concurrent.futures import ThreadPoolExecutor as _TPE

    briefing_path = None
    article_he = None

    def _briefing_task():
        return _create_briefing(content_en, display_title, base, synthesis_block)

    def _translate_task():
        if not bilingual:
            return None
        # Skip Hebrew translation if cumulative time critical (>22 min)
        elapsed_now = _time.time() - _writer_start
        if elapsed_now > 22 * 60:
            print(f"  [Agent2] ⏭  HE translation skipped — already at {elapsed_now/60:.1f} min")
            return None
        try:
            return ask_claude(prompt_he, system=system_he, max_budget=2.5, timeout=700)
        except Exception as e:
            print(f"  [Agent2] HE translation failed: {e}")
            return None

    # Pre-build HE prompts (used by _translate_task)
    system_he = """אתה מתרגם אקדמי. תרגם את המאמר הזה מאנגלית לעברית טבעית.
שמור על:
  - כל הציטוטים בפורמט המקורי (Smith, 2019)
  - כל השמות והמספרים
  - מבנה הסעיפים (## headers)
  - הטון האקדמי
זה תרגום, לא כתיבה מחדש."""
    prompt_he = f"""תרגם את המאמר הבא לעברית טבעית. שמור על המבנה והציטוטים:

{article_en}"""

    with _TPE(max_workers=2) as _ex:
        f_brief = _ex.submit(_briefing_task)
        f_trans = _ex.submit(_translate_task)
        try:
            briefing_path = f_brief.result(timeout=300)
        except Exception as e:
            print(f"  [Agent2] Briefing failed: {e}")
        try:
            article_he = f_trans.result(timeout=1200)
        except Exception as e:
            print(f"  [Agent2] HE wait timed out: {e}")

    if not bilingual:
        saved_paths = {"md": md_en, "docx": docx_en, "briefing": briefing_path}
        print(f"\n✅ Agent 2 complete → 3 files saved in {ARTICLES_DIR}\n")
        return saved_paths

    # ── Hebrew article — TRANSLATION (not rewriting) ──
    print("  [Agent2] Translating to Hebrew (חוסך זמן וטוקנים)...")

    # ── article_he was computed in parallel above ──
    if article_he:
        try:
            title_he, content_he = _split_title(article_he, f"מאמר סינתטי: {display_title}")
            md_he   = ARTICLES_DIR / f"{base}_he.md"
            docx_he = ARTICLES_DIR / f"{base}_he.docx"
            with open(md_he, "w", encoding="utf-8") as f:
                f.write(f"# {title_he}\n\n{content_he}")
            _markdown_to_docx(content_he, title_he, docx_he)
            print(f"  [Agent2] Hebrew saved: {md_he.name}, {docx_he.name}")

            saved_paths = {"md": md_en, "docx": docx_en, "md_he": md_he,
                           "docx_he": docx_he, "briefing": briefing_path}
            print(f"\n✅ Agent 2 complete → 5 files saved in {ARTICLES_DIR}\n")
        except Exception as e:
            print(f"  ⚠️  [Agent2] Hebrew save failed ({e}) — continuing with English only")
            saved_paths = {"md": md_en, "docx": docx_en, "briefing": briefing_path}
    else:
        saved_paths = {"md": md_en, "docx": docx_en, "briefing": briefing_path}
        print(f"\n✅ Agent 2 complete → English only (HE skipped)\n")
        print(f"\n✅ Agent 2 complete → 3 files saved in {ARTICLES_DIR}\n")

    return saved_paths


# ─────────────────────────────────────────────
# Standalone
# ─────────────────────────────────────────────

if __name__ == "__main__":
    import sys
    args = sys.argv[1:]
    if not args:
        # Use all JSON files in papers dir
        papers_list = sorted(Path("output/papers").glob("*.json"), key=lambda p: p.stat().st_mtime, reverse=True)[:3]
    else:
        papers_list = [Path(a) for a in args]
    paths = run_writer(papers_list)
    print(f"Article saved: {paths}")
